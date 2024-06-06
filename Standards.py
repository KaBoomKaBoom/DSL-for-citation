from docx import Document
from docx.shared import Pt
import datetime

class Standards:
    def generate_apa_citation(qouote, source_type, author, year, title, book_title, publisher):
        #print(qouote, source_type, author, year, title, book_title, publisher)
        doc = Document("test.docx")
            # For in-text citation
        citation_text = qouote + "(" + author + ", " + str(year) + ")"
        run = doc.add_paragraph()
        run = run.add_run(citation_text)
        font = run.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
            # For bibliography/reference list
        citation_bib = author + " (" + str(year) + "). " + title + ". "


        if source_type == "book":
            citation_bib += book_title  + ' ' + publisher + "."
            run = doc.add_paragraph
            run = run.add_run(citation_bib)
            font = run.font
            font.name = 'Times New Roman'
            font.size = Pt(12)
        elif source_type == "journal":
            citation_bib += book_title  + ' ' + publisher + "."
            run = doc.add_paragraph
            run = run.add_run(citation_bib)
            font = run.font
            font.name = 'Times New Roman'
            font.size = Pt(12)
        elif source_type == "website":
            run = doc.add_paragraph
            run = run.add_run(citation_bib)
            font = run.font
            font.name = 'Times New Roman'
            font.size = Pt(12)
        doc.save('test.docx')
        return citation_text, citation_bib

    def generate_mla_citation(qouote, source_type, author, year, title, book_title, publisher):
        #print(qouote, source_type, author, year, title, book_title, publisher)
        doc = Document("test.docx")
            # For in-text citation
        citation_text = f"{qouote} ({author})."
        run = doc.add_paragraph()
        run = run.add_run(citation_text)
        font = run.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
            # For bibliography/reference list
        citation_bib = author + '"' + title + '."'
        

        if source_type == "book":
            citation_bib += book_title  + ', ' + publisher + ', ' + str(year) +  "."
            run = doc.add_paragraph()
            run = run.add_run(f'{author} \"{title}.\"{book_title}, {publisher}, {str(year)}.')
            font = run.font
            font.name = 'Times New Roman'
            font.size = Pt(12)
        elif source_type == "journal":
            citation_bib += book_title  + ', ' + publisher + ', ' + str(year) +  "."
            run = doc.add_paragraph()
            run = run.add_run(f'{author} \"{title}.\"{book_title}, {publisher}, {str(year)}.')
            font = run.font
            font.name = 'Times New Roman'
            font.size = Pt(12)
            
        elif source_type == "website":
            citation_bib += book_title  + ', ' + publisher + ', ' + str(year) +  "."
            run = doc.add_paragraph()
            run = run.add_run(f'{author} \"{title}.\"{book_title}, {publisher}, {str(year)}.')
            font = run.font
            font.name = 'Times New Roman'
            font.size = Pt(12)
        doc.save('test.docx')
        return citation_text, citation_bib

    def generate_cms_citation(qouote, source_type, author, year, title, book_title, publisher, link):
        #print(qouote, source_type, author, year, title, book_title, publisher)
        doc = Document("test.docx")
            # For in-text citation
        citation_text = f"{qouote} ({author} {str(year)})."
        run = doc.add_paragraph()
        run = run.add_run(citation_text)
        font = run.font
        font.name = 'Times New Roman'
        font.size = Pt(12)

            # For bibliography/reference list
        citation_bib = author + ". "

        if source_type == "book":
            citation_bib += book_title  + ', ' + publisher + ', ' + str(year) +  "."
            run = doc.add_paragraph()
            run = run.add_run(f"{author}. \"{book_title},\" {publisher}, {str(year)}.")
            font = run.font
            font.name = 'Times New Roman'
            font.size = Pt(12)

        elif source_type == "journal":
            citation_bib += '"' + book_title  + '," ' + publisher + ', ' + str(year) +  "."
            run = doc.add_paragraph()
            run = run.add_run(f"{author}. \"{book_title},\" {publisher}, {str(year)}.")
            font = run.font
            font.name = 'Times New Roman'
            font.size = Pt(12)

        elif source_type == "website":
            date = datetime.datetime.now()
            citation_bib += title  + ', ' + publisher + ', ' + 'accessed ' + date.strftime("%d %b %Y") + ", " +  link + "."
            run = doc.add_paragraph()
            run = run.add_run(f"{author}. {title}, {publisher}, accessed {date.strftime('%d %b %Y')}, {link}.")
            font = run.font
            font.name = 'Times New Roman'
            font.size = Pt(12)

        doc.save('test.docx')
        return citation_text, citation_bib

    def generate_cse_citation(qouote, source_type, author, year, title, book_title, publisher, link, index):
        #print(qouote, source_type, author, year, title, book_title, publisher)
        doc = Document("test.docx")
            # For in-text citation superscript the number of the citation
        citation_text = qouote + ' ' + str(index) + "."
        run = doc.add_paragraph()
        run = run.add_run(qouote + "[" + str(index) + "]")
        font = run.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
            # For bibliography/reference list
        citation_bib = author  

        if source_type == "book":
            citation_bib += book_title  + ', ' + publisher + ', ' + str(year) +  "."
            run = doc.add_paragraph()
            run = run.add_run(f"{author}{book_title}. {publisher}; {str(year)}.")
            font = run.font
            font.name = 'Times New Roman'
            font.size = Pt(12)

        elif source_type == "journal":
            citation_bib += '"' + book_title  + '," ' + publisher + ', ' + str(year) +  "."
            run = doc.add_paragraph()
            run = run.add_run(f"{author} {book_title}. {publisher}. {str(year)}")
            font = run.font
            font.name = 'Times New Roman'
            font.size = Pt(12)

        elif source_type == "website":
            date = datetime.datetime.now()
            citation_bib += title  + ', ' + publisher + '. ' + str(year) +', ' + '[accessed ' + date.strftime("%d %b %Y") + "]; " +  link + "."
            run = doc.add_paragraph()
            run = run.add_run(f"{author} {title}. {publisher}. {str(year)} [accesed {date.strftime("%d %b %Y")}]; {link}.")
            font = run.font
            font.name = 'Times New Roman'
            font.size = Pt(12)
        doc.save('test.docx')
        return citation_text, citation_bib

    def cite_iso(qouote, source_type, author, year,  title, book_title,  publisher, link, index):
        #print(qouote, source_type, author, year, title, book_title, publisher)
        doc = Document("test.docx")
            # For in-text citation superscript the number of the citation
        citation_text = qouote + ' [' + str(index) + "]."
        run = doc.add_paragraph()
        run = run.add_run(citation_text)
        font = run.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
            # For bibliography/reference list
        citation_bib = f'[{str(index)}]  ' + author + '. ' 

        if source_type == "book":
            citation_bib += book_title  + '. ' + publisher + ', ' + str(year) +  "."
            run = doc.add_paragraph()
            run = run.add_run(f"[{str(index)}]  {author}. {book_title}. {publisher}; {str(year)}.")
            font = run.font
            font.name = 'Times New Roman'
            font.size = Pt(12)

        elif source_type == "journal":
            citation_bib += '\"' + book_title  + ',\" ' + publisher + ', ' + str(year) +  "."
            run = doc.add_paragraph()
            run = run.add_run(f"[{str(index)}]  {author}. {book_title}. {publisher}; {str(year)}.")
            font = run.font
            font.name = 'Times New Roman'
            font.size = Pt(12)

        elif source_type == "website":
            date = datetime.datetime.now()
            citation_bib += f'{title}. {publisher}. {str(year)}. Available from: {link}. Accessed Date {date.strftime("%d %b %Y")}.'
            run = doc.add_paragraph()
            run = run.add_run(f"[{str(index)}]  {author}, \"{title}\". {publisher}. {str(year)} Available from: {link}. Accessed Date {date.strftime("%d %b %Y")}.")
            font = run.font
            font.name = 'Times New Roman'
            font.size = Pt(12)
        doc.save('test.docx')
        return citation_text, citation_bib
    
    def cite_ieee(qouote, source_type, author, year,  title, book_title,  publisher, link, index = None):
        doc = Document("test.docx")
            # For in-text citation superscript the number of the citation
        citation_text = qouote + ' [' + str(index) + "]."
        run = doc.add_paragraph()
        run = run.add_run(qouote + "[" + str(index) + "]")
        font = run.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
            # For bibliography/reference list
        citation_bib = author + '. ' 

        if source_type == "book":
            citation_bib += book_title  + '. ' + publisher + ', ' + str(year) +  "."
            run = doc.add_paragraph()
            run = run.add_run(f"[{str(index)}]  {author}. {book_title}. {publisher}; {str(year)}.")
            font = run.font
            font.name = 'Times New Roman'
            font.size = Pt(12)

        elif source_type == "journal":
            citation_bib += '"' + book_title  + '," ' + publisher + ', ' + str(year) +  "."
            run = doc.add_paragraph()
            run = run.add_run(f"[{str(index)}]  {author}. {book_title}. {publisher}; {str(year)}.")
            font = run.font
            font.name = 'Times New Roman'
            font.size = Pt(12)

        elif source_type == "website":
            date = datetime.datetime.now()
            citation_bib += title  + ', ' + publisher + '. ' + str(year) +', ' + '[accessed ' + date.strftime("%d %b %Y") + "]; " +  link + "."
            run = doc.add_paragraph()
            run = run.add_run(f"[{str(index)}]  {author}, {title}. {publisher}. {str(year)} Available: {link} [Accessed Date {date.strftime("%d %b %Y")}].")
            font = run.font
            font.name = 'Times New Roman'
            font.size = Pt(12)
        doc.save('test.docx')
        return citation_text, citation_bib